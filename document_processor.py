from __future__ import annotations

from pathlib import Path
from typing import Callable, Iterable

import pymupdf
from docx import Document
from docx.enum.text import WD_BREAK
from docx.table import Table
from docx.text.paragraph import Paragraph

from translator_engine import OnlineTranslator


class DocumentProcessor:
    def __init__(
        self,
        translator: OnlineTranslator,
        progress_callback: Callable[[int, int], None] | None = None,
        log_callback: Callable[[str], None] | None = None,
    ) -> None:
        self.translator = translator
        self.progress = progress_callback or (lambda _done, _total: None)
        self.log = log_callback or (lambda _msg: None)

    @staticmethod
    def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)

    @staticmethod
    def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                yield from cell.paragraphs
                for nested in cell.tables:
                    yield from DocumentProcessor._iter_table_paragraphs(nested)

    def _collect_docx_paragraphs(self, document, translate_tables: bool, translate_headers: bool) -> list[Paragraph]:
        paragraphs: list[Paragraph] = list(document.paragraphs)
        if translate_tables:
            for table in document.tables:
                paragraphs.extend(self._iter_table_paragraphs(table))
        if translate_headers:
            seen_parts: set[str] = set()
            for section in document.sections:
                for container in [
                    section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer,
                ]:
                    partname = str(container.part.partname)
                    if partname in seen_parts:
                        continue
                    seen_parts.add(partname)
                    paragraphs.extend(container.paragraphs)
                    if translate_tables:
                        for table in container.tables:
                            paragraphs.extend(self._iter_table_paragraphs(table))
        return paragraphs

    def _process_docx(self, source_path: Path, output_path: Path, translate_tables: bool, translate_headers: bool) -> Path:
        self.log("Открытие DOCX…")
        document = Document(str(source_path))
        paragraphs = self._collect_docx_paragraphs(document, translate_tables, translate_headers)
        targets = [p for p in paragraphs if self.translator.should_translate(p.text)]
        total = len(targets)
        self.log(f"Блоков для перевода: {total}")
        self.progress(0, total)

        for index, paragraph in enumerate(targets, start=1):
            self._replace_paragraph_text(paragraph, self.translator.translate_text(paragraph.text))
            if index == 1 or index % 5 == 0 or index == total:
                self.progress(index, total)
                self.log(f"Переведено: {index}/{total}")

        document.save(str(output_path))
        return output_path

    def _process_pdf(self, source_path: Path, output_path: Path) -> Path:
        self.log("Извлечение текста из PDF…")
        pdf = pymupdf.open(str(source_path))
        pages = []
        for page in pdf:
            blocks = page.get_text("blocks", sort=True)
            pages.append([str(block[4]).strip() for block in blocks if str(block[4]).strip()])

        targets = [text for page in pages for text in page if self.translator.should_translate(text)]
        total = len(targets)
        self.log(f"Блоков для перевода: {total}")
        self.progress(0, total)

        output = Document()
        done = 0
        for page_number, page_blocks in enumerate(pages, start=1):
            output.add_heading(f"Страница {page_number}", level=2)
            for block_text in page_blocks:
                output.add_paragraph(self.translator.translate_text(block_text))
                if self.translator.should_translate(block_text):
                    done += 1
                    if done == 1 or done % 5 == 0 or done == total:
                        self.progress(done, total)
                        self.log(f"Переведено: {done}/{total}")
            if page_number != len(pages):
                output.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        output.save(str(output_path))
        pdf.close()
        return output_path

    def process(self, source_path: Path, output_dir: Path, translate_tables: bool = True, translate_headers: bool = True) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{source_path.stem}_RU.docx"
        if source_path.suffix.lower() == ".docx":
            return self._process_docx(source_path, output_path, translate_tables, translate_headers)
        if source_path.suffix.lower() == ".pdf":
            return self._process_pdf(source_path, output_path)
        raise ValueError("Поддерживаются только DOCX и PDF.")
