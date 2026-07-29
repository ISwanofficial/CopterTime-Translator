from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from coptertime_docs.storage.database import Database
from coptertime_docs.translator.engine import TranslationEngine


ProgressCallback = Callable[[int, int, str], None]


@dataclass
class DocumentAnalysis:
    paragraphs: int
    text_fragments: int
    tables: int
    images: int
    words: int
    english_fragments: int


@dataclass
class ProcessingReport:
    total_fragments: int
    translated_new: int
    translated_from_memory: int
    skipped: int
    english_fragments_after: int
    output_path: Path


class DocxProcessor:
    def __init__(self, database: Database, engine: TranslationEngine) -> None:
        self.database = database
        self.engine = engine

    def analyze(self, path: Path) -> DocumentAnalysis:
        document = Document(path)
        paragraphs = list(self._iter_paragraphs(document))
        text_fragments = [p.text.strip() for p in paragraphs if p.text.strip()]
        words = sum(len(text.split()) for text in text_fragments)
        english = sum(1 for text in text_fragments if self._contains_english(text))
        images = len(document.part._package.image_parts)

        return DocumentAnalysis(
            paragraphs=len(paragraphs),
            text_fragments=len(text_fragments),
            tables=len(document.tables),
            images=images,
            words=words,
            english_fragments=english,
        )

    def translate(
        self,
        source_path: Path,
        output_path: Path,
        brand: str,
        progress: ProgressCallback | None = None,
    ) -> ProcessingReport:
        source_path = Path(source_path)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        shutil.copy2(source_path, output_path)
        document = Document(output_path)

        paragraphs = [
            p for p in self._iter_paragraphs(document)
            if p.text.strip() and self._contains_english(p.text)
        ]
        total = len(paragraphs)
        translated_new = 0
        translated_from_memory = 0
        skipped = 0

        for index, paragraph in enumerate(paragraphs, start=1):
            source_text = paragraph.text.strip()
            if progress:
                progress(index, total, source_text[:90])

            try:
                result = self.engine.translate(source_text, brand)
                if result.text and result.text != source_text:
                    self._replace_paragraph_text(paragraph, result.text)
                    if result.from_memory:
                        translated_from_memory += 1
                    else:
                        translated_new += 1
                else:
                    skipped += 1
            except Exception:
                # Сохраняем исходный фрагмент, чтобы один сбой не уничтожил документ.
                skipped += 1

        document.save(output_path)

        english_after = self.analyze(output_path).english_fragments
        status = "completed" if skipped == 0 else "completed_with_warnings"
        self.database.record_document(
            str(source_path),
            str(output_path),
            brand,
            status,
            total,
            translated_new,
            translated_from_memory,
            english_after,
        )

        return ProcessingReport(
            total_fragments=total,
            translated_new=translated_new,
            translated_from_memory=translated_from_memory,
            skipped=skipped,
            english_fragments_after=english_after,
            output_path=output_path,
        )

    @staticmethod
    def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(text)
            return

        first_run = paragraph.runs[0]
        first_run.text = text
        for run in paragraph.runs[1:]:
            run.text = ""

    @classmethod
    def _iter_paragraphs(
        cls,
        parent: DocumentObject | _Cell,
    ) -> Iterable[Paragraph]:
        for paragraph in parent.paragraphs:
            yield paragraph

        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cls._iter_paragraphs(cell)

        if isinstance(parent, DocumentObject):
            for section in parent.sections:
                for paragraph in section.header.paragraphs:
                    yield paragraph
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cls._iter_paragraphs(cell)

                for paragraph in section.footer.paragraphs:
                    yield paragraph
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cls._iter_paragraphs(cell)

    @staticmethod
    def _contains_english(text: str) -> bool:
        latin = len(re.findall(r"[A-Za-z]", text or ""))
        cyrillic = len(re.findall(r"[А-Яа-яЁё]", text or ""))
        return latin >= 2 and latin > cyrillic * 0.35
